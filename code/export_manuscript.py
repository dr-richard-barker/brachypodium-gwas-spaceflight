#!/usr/bin/env python3
"""
Convert the Brachypodium GWAS-Spaceflight manuscript from LaTeX to PDF and Word.

Parses the manuscript.tex file, strips LaTeX commands, and generates:
  - manuscript/manuscript.pdf (via fpdf2)
  - manuscript/manuscript.docx (via python-docx)
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from fpdf import FPDF


MANUSCRIPT_DIR = Path(__file__).resolve().parent.parent / "manuscript"
TEX_FILE = MANUSCRIPT_DIR / "manuscript.tex"


# =============================================================================
# LaTeX Parser
# =============================================================================

def parse_latex(tex_path: Path) -> dict:
    """Parse LaTeX manuscript into structured sections."""
    raw = tex_path.read_text(encoding="utf-8")

    # Extract metadata
    title_m = re.search(r"\\title\[.*?\]\{(.+?)\}", raw, re.DOTALL)
    title = clean_latex(title_m.group(1)) if title_m else "Untitled"

    author_m = re.search(r"\\author\*?\[.*?\]\{\\fnm\{(.+?)\}\s*\\sur\{(.+?)\}\}", raw)
    author = f"{author_m.group(1)} {author_m.group(2)}" if author_m else "Unknown"

    email_m = re.search(r"\\email\{(.+?)\}", raw)
    email = email_m.group(1) if email_m else ""

    affil_m = re.search(r"\\affil\[.*?\]\{\\orgname\{(.+?)\}\}", raw)
    affiliation = affil_m.group(1) if affil_m else ""

    abstract_m = re.search(r"\\abstract\{(.+?)\}\s*\n", raw, re.DOTALL)
    abstract = clean_latex(abstract_m.group(1)) if abstract_m else ""

    keywords_m = re.search(r"\\keywords\{(.+?)\}", raw, re.DOTALL)
    keywords = clean_latex(keywords_m.group(1)) if keywords_m else ""

    # Extract body (between \mainmatter and \bibliographystyle or \end{document})
    body_m = re.search(
        r"\\mainmatter\s*(.+?)(?:\\bibliographystyle|\\end\{document\})",
        raw, re.DOTALL
    )
    body_tex = body_m.group(1) if body_m else ""

    # Parse sections
    sections = parse_sections(body_tex)

    # Extract acknowledgments and author contributions
    ack_m = re.search(r"\\section\*\{Acknowledgments\}\s*(.+?)(?=\\section|$)", raw, re.DOTALL)
    ack = clean_latex(ack_m.group(1)) if ack_m else ""

    contrib_m = re.search(r"\\section\*\{Author contributions\}\s*(.+?)(?=\\section|$)", raw, re.DOTALL)
    contrib = clean_latex(contrib_m.group(1)) if contrib_m else ""

    competing_m = re.search(r"\\section\*\{Competing interests\}\s*(.+?)(?=\\section|\\biblio|$)", raw, re.DOTALL)
    competing = clean_latex(competing_m.group(1)) if competing_m else ""

    return {
        "title": title,
        "author": author,
        "email": email,
        "affiliation": affiliation,
        "abstract": abstract,
        "keywords": keywords,
        "sections": sections,
        "acknowledgments": ack,
        "author_contributions": contrib,
        "competing_interests": competing,
    }


def parse_sections(body_tex: str) -> list:
    """Parse LaTeX body into a list of (level, title, content) tuples."""
    # Split on \section, \subsection, \subsection*
    pattern = r"\\(section|subsection)\*?\{(.+?)\}"
    parts = re.split(pattern, body_tex)

    sections = []
    i = 0
    while i < len(parts):
        if i == 0:
            # Content before first section
            text = clean_latex(parts[0].strip())
            if text:
                sections.append((0, "", text))
            i += 1
        elif parts[i] in ("section", "subsection"):
            level = 1 if parts[i] == "section" else 2
            title = clean_latex(parts[i + 1])
            # Content is in parts[i+2] (up to next section marker)
            content = clean_latex(parts[i + 2].strip()) if i + 2 < len(parts) else ""
            sections.append((level, title, content))
            i += 3
        else:
            i += 1

    return sections


def clean_latex(text: str) -> str:
    """Strip LaTeX commands and clean up text for plain output."""
    # Remove comments
    text = re.sub(r"(?<!\\)%.*$", "", text, flags=re.MULTILINE)

    # Handle common commands
    text = re.sub(r"\\textit\{(.+?)\}", r"\1", text)
    text = re.sub(r"\\textbf\{(.+?)\}", r"\1", text)
    text = re.sub(r"\\emph\{(.+?)\}", r"\1", text)
    text = re.sub(r"\\texttt\{(.+?)\}", r"\1", text)
    text = re.sub(r"\\text\{(.+?)\}", r"\1", text)

    # Math
    text = re.sub(r"\$(.+?)\$", r"\1", text)
    text = re.sub(r"\\times", "×", text)
    text = re.sub(r"\\sim", "~", text)
    text = re.sub(r"\\rho", "ρ", text)
    text = re.sub(r"\\beta", "β", text)
    text = re.sub(r"\\log_?2?", "log₂", text)
    text = re.sub(r"\\geq?", "≥", text)
    text = re.sub(r"\\leq?", "≤", text)
    text = re.sub(r"\\infty", "∞", text)
    text = re.sub(r"\^\{?(-?\d+)\}?", lambda m: _superscript(m.group(1)), text)
    text = re.sub(r"_\{?([a-zA-Z0-9]+)\}?", lambda m: _subscript(m.group(1)), text)

    # URLs
    text = re.sub(r"\\url\{(.+?)\}", r"\1", text)
    text = re.sub(r"\\href\{(.+?)\}\{(.+?)\}", r"\2 (\1)", text)

    # Citations
    text = re.sub(r"~?\\cite\{(.+?)\}", lambda m: f"[{m.group(1)}]", text)

    # Cross-references
    text = re.sub(r"\\ref\{(.+?)\}", r"[\1]", text)
    text = re.sub(r"\\label\{(.+?)\}", "", text)

    # Remove remaining commands
    text = re.sub(r"\\textdegree", "°", text)
    text = re.sub(r"\\textsuperscript\{(.+?)\}", r"\1", text)
    text = re.sub(r"\\[a-zA-Z]+\{(.+?)\}", r"\1", text)
    text = re.sub(r"\\[a-zA-Z]+", " ", text)
    text = re.sub(r"[{}]", "", text)

    # Clean whitespace
    text = re.sub(r"\n\s*\n\s*\n+", "\n\n", text)
    text = re.sub(r"  +", " ", text)
    text = text.strip()

    return text


def _superscript(s: str) -> str:
    sup = str.maketrans("0123456789-", "⁰¹²³⁴⁵⁶⁷⁸⁹⁻")
    return s.translate(sup)


def _subscript(s: str) -> str:
    sub = str.maketrans("0123456789aehijklmnoprstuvx", "₀₁₂₃₄₅₆₇₈₉ₐₑₕᵢⱼₖₗₘₙₒₚᵣₛₜᵤᵥₓ")
    try:
        return s.translate(sub)
    except Exception:
        return s


# =============================================================================
# Word Document Generator
# =============================================================================

def generate_docx(parsed: dict, output_path: Path):
    """Generate a formatted Word document from parsed LaTeX."""
    doc = Document()

    # --- Styles ---
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(parsed["title"])
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = RGBColor(0x0B, 0x1D, 0x3A)

    # Author line
    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_run = author_para.add_run(
        f"{parsed['author']}\n{parsed['affiliation']}\n{parsed['email']}"
    )
    author_run.font.size = Pt(11)
    author_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Draft watermark note
    draft_para = doc.add_paragraph()
    draft_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    draft_run = draft_para.add_run("— DRAFT MANUSCRIPT —")
    draft_run.bold = True
    draft_run.font.size = Pt(12)
    draft_run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    doc.add_paragraph("")  # spacer

    # Abstract
    doc.add_heading("Abstract", level=1)
    abstract_para = doc.add_paragraph(parsed["abstract"])
    abstract_para.paragraph_format.first_line_indent = Pt(0)
    for run in abstract_para.runs:
        run.font.size = Pt(10)

    # Keywords
    kw_para = doc.add_paragraph()
    kw_run = kw_para.add_run("Keywords: ")
    kw_run.bold = True
    kw_run.font.size = Pt(10)
    kw_para.add_run(parsed["keywords"]).font.size = Pt(10)

    doc.add_page_break()

    # Sections
    for level, title, content in parsed["sections"]:
        if title:
            doc.add_heading(title, level=min(level, 3))
        if content:
            # Split into paragraphs
            for para_text in content.split("\n\n"):
                para_text = para_text.strip()
                if para_text:
                    # Check for TODO/PLACEHOLDER markers
                    if "TODO" in para_text or "PLACEHOLDER" in para_text:
                        p = doc.add_paragraph()
                        run = p.add_run(para_text)
                        run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)
                        run.italic = True
                    else:
                        doc.add_paragraph(para_text)

    # Back matter
    if parsed["acknowledgments"]:
        doc.add_heading("Acknowledgments", level=1)
        doc.add_paragraph(parsed["acknowledgments"])

    if parsed["author_contributions"]:
        doc.add_heading("Author Contributions", level=1)
        doc.add_paragraph(parsed["author_contributions"])

    if parsed["competing_interests"]:
        doc.add_heading("Competing Interests", level=1)
        doc.add_paragraph(parsed["competing_interests"])

    # Save
    doc.save(str(output_path))
    print(f"  ✓ Word document saved: {output_path}")


# =============================================================================
# PDF Generator
# =============================================================================

class ManuscriptPDF(FPDF):
    """Custom PDF class for the manuscript."""

    def __init__(self):
        super().__init__()
        self.set_auto_page_break(auto=True, margin=25)

    def header(self):
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(150, 150, 150)
        self.cell(0, 8, "DRAFT -- Brachypodium GWAS-Spaceflight Integration", align="C")
        self.ln(10)

    def footer(self):
        self.set_y(-15)
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(150, 150, 150)
        self.cell(0, 10, f"Page {self.page_no()}/{{nb}}", align="C")

    def chapter_title(self, title: str, level: int = 1):
        sizes = {1: 14, 2: 12, 3: 11}
        size = sizes.get(level, 11)
        self.set_font("Helvetica", "B", size)
        self.set_text_color(11, 29, 58)  # Navy
        self.ln(4 if level > 1 else 8)
        safe_title = sanitize_pdf_text(title)
        self.multi_cell(0, 6, safe_title)
        self.ln(3)

    def body_text(self, text: str, italic: bool = False, color: tuple = (0, 0, 0)):
        self.set_font("Helvetica", "I" if italic else "", 10)
        self.set_text_color(*color)
        safe_text = sanitize_pdf_text(text)
        self.multi_cell(0, 5, safe_text)
        self.ln(2)


def sanitize_pdf_text(text: str) -> str:
    """Sanitize unicode characters for standard PDF Helvetica font."""
    replacements = {
        "—": "--",
        "–": "-",
        "“": '"',
        "”": '"',
        "‘": "'",
        "’": "'",
        "°": " deg",
        "±": "+/-",
        "×": "x",
        "α": "alpha",
        "β": "beta",
        "ρ": "rho",
        "μ": "micro",
        "θ": "theta",
        "…": "...",
        "≥": ">=",
        "≤": "<=",
        "²": "^2",
        "³": "^3",
        "¹": "^1",
        "⁰": "^0",
        "₄": "_4",
        "₀": "_0",
        "•": "*",
        "→": "->",
    }
    for k, v in replacements.items():
        text = text.replace(k, v)
    return text.encode("latin-1", errors="replace").decode("latin-1")


def generate_pdf(parsed: dict, output_path: Path):
    """Generate a formatted PDF from parsed LaTeX."""
    pdf = ManuscriptPDF()
    pdf.alias_nb_pages()
    pdf.add_page()

    # Title
    pdf.set_font("Helvetica", "B", 18)
    pdf.set_text_color(11, 29, 58)
    title_safe = sanitize_pdf_text(parsed["title"])
    pdf.multi_cell(0, 8, title_safe, align="C")
    pdf.ln(5)

    # Author
    pdf.set_font("Helvetica", "", 11)
    pdf.set_text_color(80, 80, 80)
    pdf.cell(0, 6, sanitize_pdf_text(parsed["author"]), align="C")
    pdf.ln(5)
    pdf.set_font("Helvetica", "I", 10)
    pdf.cell(0, 5, sanitize_pdf_text(parsed["affiliation"]), align="C")
    pdf.ln(4)
    pdf.cell(0, 5, sanitize_pdf_text(parsed["email"]), align="C")
    pdf.ln(8)

    # DRAFT notice
    pdf.set_font("Helvetica", "B", 12)
    pdf.set_text_color(204, 0, 0)
    pdf.cell(0, 8, "-- DRAFT MANUSCRIPT --", align="C")
    pdf.ln(10)

    # Abstract
    pdf.chapter_title("Abstract", 1)
    pdf.body_text(parsed["abstract"])

    # Keywords
    pdf.set_font("Helvetica", "B", 10)
    pdf.set_text_color(0, 0, 0)
    pdf.cell(20, 5, "Keywords: ")
    pdf.set_font("Helvetica", "I", 10)
    kw_safe = sanitize_pdf_text(parsed["keywords"])
    pdf.multi_cell(0, 5, kw_safe)
    pdf.ln(5)

    # Sections
    for level, title, content in parsed["sections"]:
        if title:
            pdf.chapter_title(title, level)
        if content:
            for para_text in content.split("\n\n"):
                para_text = para_text.strip()
                if not para_text:
                    continue
                if "TODO" in para_text or "PLACEHOLDER" in para_text:
                    pdf.body_text(para_text, italic=True, color=(204, 102, 0))
                else:
                    pdf.body_text(para_text)

    # Back matter
    if parsed["acknowledgments"]:
        pdf.chapter_title("Acknowledgments", 1)
        pdf.body_text(parsed["acknowledgments"])

    if parsed["author_contributions"]:
        pdf.chapter_title("Author Contributions", 1)
        pdf.body_text(parsed["author_contributions"])

    if parsed["competing_interests"]:
        pdf.chapter_title("Competing Interests", 1)
        pdf.body_text(parsed["competing_interests"])

    # Save
    pdf.output(str(output_path))
    print(f"  ✓ PDF saved: {output_path}")


# =============================================================================
# Main
# =============================================================================

def main():
    print("=" * 60)
    print("Manuscript Export: LaTeX → PDF + Word")
    print("=" * 60)

    if not TEX_FILE.exists():
        print(f"ERROR: {TEX_FILE} not found")
        sys.exit(1)

    print(f"\nParsing: {TEX_FILE}")
    parsed = parse_latex(TEX_FILE)

    print(f"  Title:  {parsed['title'][:60]}...")
    print(f"  Author: {parsed['author']}")
    n_sections = len(parsed["sections"])
    print(f"  Sections: {n_sections}")

    # Generate outputs
    pdf_path = MANUSCRIPT_DIR / "manuscript.pdf"
    docx_path = MANUSCRIPT_DIR / "manuscript.docx"

    print("\nGenerating documents...")
    generate_pdf(parsed, pdf_path)
    generate_docx(parsed, docx_path)

    print(f"\n✓ Both documents saved to {MANUSCRIPT_DIR}/")
    print(f"  PDF:  {pdf_path.stat().st_size / 1024:.1f} KB")
    print(f"  DOCX: {docx_path.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
